"""
build_formulas_doc.py
Собирает Word-документ со всеми формулами из DST_ALGORITHM_FULL.md.

Структура документа:
  - Заголовок документа
  - Таблица обозначений (§2 целиком)
  - Для каждого раздела: заголовок раздела, затем все формулы
    (строки-цитаты ">" с математическим содержимым и блоки $$)
    каждая в отдельном абзаце моноширинным шрифтом с рамкой

Зависимости: python-docx
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_FILE = Path(__file__).parent / "DST_ALGORITHM_FULL.md"
OUT_FILE = Path(__file__).parent / "DST_formulas.docx"

# ── Вспомогательные функции ──────────────────────────────────────────────────

def add_heading(doc, text, level):
    """Добавляет заголовок нужного уровня."""
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def set_cell_border(cell, **kwargs):
    """Устанавливает границы ячейки таблицы."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), kwargs.get(edge, 'single'))
        tag.set(qn('w:sz'), '4')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), '888888')
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_formula_block(doc, formula_text):
    """
    Добавляет формулу в документ в виде абзаца с серым фоном
    и моноширинным шрифтом.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Серый фон абзаца через shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)

    # Отступы
    pf = OxmlElement('w:pBdr')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), '4')
        tag.set(qn('w:space'), '4' if edge in ('left', 'right') else '2')
        tag.set(qn('w:color'), 'AAAAAA')
        pf.append(tag)
    pPr.append(pf)

    run = p.add_run(formula_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    return p


def add_normal_para(doc, text):
    """Добавляет обычный текстовый абзац."""
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    return p


def is_formula_line(line):
    """
    Возвращает True, если строка — формула (цитата с математическим содержимым).
    Цитаты-примечания (начинаются с > ** или содержат только текст) не считаем формулами.
    """
    if not line.startswith('>'):
        return False
    content = line.lstrip('> ').strip()
    # Пропускаем пустые цитаты
    if not content:
        return False
    # Пропускаем чисто текстовые примечания (не содержат математики)
    math_markers = [
        '=', '≠', '∅', '∑', '∈', '∉', '∪', '∩', '∀', '∃',
        '≤', '≥', '·', '−', '/', '|{', '}|', '←', ':=',
        'κ', 'λ', 'σ', 'μ', 'δ', 'ρ', 'γ', 'π', 'ω', 'Δ',
        'm(', 'm′', 'm*', 'BetP', 'num(', 'norm', 'BPA',
        '+= ', '+=min',
    ]
    has_math = any(m in content for m in math_markers)
    # Пропускаем строки-примечания вроде "> **Почему..."
    is_note = content.startswith('**') or content.startswith('Ключевое') or \
              content.startswith('В текущих') or content.startswith('Если ')
    return has_math and not is_note


def strip_bold(text):
    """Убирает markdown-**bold** маркеры."""
    return re.sub(r'\*\*(.+?)\*\*', r'\1', text)


# ── Парсинг Markdown ─────────────────────────────────────────────────────────

def parse_md(path):
    """
    Возвращает список секций:
      {'heading': str, 'level': int, 'formulas': [str], 'table_rows': [...]}
    Секция с заголовком "Таблица обозначений" дополнительно содержит table_rows.
    """
    lines = path.read_text(encoding='utf-8').splitlines()

    sections = []
    current = None
    in_latex = False
    latex_buf = []
    in_table = False

    def flush_section():
        if current is not None:
            sections.append(current)

    for line in lines:
        # Заголовок
        m = re.match(r'^(#{1,3})\s+(.+)', line)
        if m:
            flush_section()
            level = len(m.group(1))
            heading = m.group(2).strip()
            current = {
                'heading': heading,
                'level': level,
                'formulas': [],
                'table_rows': [],
                'context': [],   # строки контекста между формулами
            }
            in_table = False
            continue

        if current is None:
            continue

        # LaTeX-блок $$
        if line.strip() == '$$':
            if not in_latex:
                in_latex = True
                latex_buf = []
            else:
                in_latex = False
                formula = ' '.join(latex_buf).strip()
                if formula:
                    current['formulas'].append(('latex', formula))
            continue

        if in_latex:
            latex_buf.append(line.strip())
            continue

        # Таблица обозначений
        if '|' in line and 'Таблица обозначений' in current.get('heading', ''):
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if cells and cells[0] not in ('Символ', '------', '-----', '----'):
                if len(cells) >= 3:
                    current['table_rows'].append(cells[:3])
            continue

        # Формула-цитата
        if is_formula_line(line):
            content = re.sub(r'^>\s*', '', line).strip()
            content = strip_bold(content)
            current['formulas'].append(('inline', content))
            continue

        # Строка контекста (заголовок подраздела, пояснение перед формулой)
        if line.startswith('**') and line.endswith('**'):
            current['context'].append(line.strip('*').strip())
        elif re.match(r'^\*[^*]', line) and line.endswith('*'):
            # Курсивная строка — подзаголовок уровня (Уровень 0, Уровень 1…)
            current['context'].append(line.strip('*').strip())

    flush_section()
    return sections


# ── Сборка Word-документа ────────────────────────────────────────────────────

def build_doc(sections, out_path):
    doc = Document()

    # Поля страницы
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    # Заголовок документа
    title = doc.add_heading('Формулы DST-алгоритма JAcademicSupport', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        'Данный документ содержит все математические формулы из файла '
        'DST_ALGORITHM_FULL.md. Формулы сгруппированы по разделам.'
    )

    for sec in sections:
        heading = sec['heading']
        level = min(sec['level'], 3)

        # Раздел таблицы обозначений — отдельная обработка
        if 'Таблица обозначений' in heading and sec['table_rows']:
            add_heading(doc, heading, level)
            rows = sec['table_rows']
            tbl = doc.add_table(rows=1 + len(rows), cols=3)
            tbl.style = 'Table Grid'
            # Заголовок таблицы
            hdr = tbl.rows[0].cells
            for i, h in enumerate(['Символ', 'Тип', 'Описание']):
                hdr[i].text = h
                run = hdr[i].paragraphs[0].runs[0]
                run.bold = True
                run.font.size = Pt(9)
            # Данные
            for ri, row_data in enumerate(rows):
                row = tbl.rows[ri + 1].cells
                for ci, val in enumerate(row_data[:3]):
                    row[ci].text = val
                    row[ci].paragraphs[0].runs[0].font.size = Pt(9)
            doc.add_paragraph()
            continue

        # Разделы без формул — пропускаем (не добавляем пустые)
        if not sec['formulas']:
            continue

        add_heading(doc, heading, level)

        # Выводим формулы, при необходимости — с контекстными подписями
        formulas = sec['formulas']
        for ftype, ftext in formulas:
            add_formula_block(doc, ftext)

        doc.add_paragraph()   # пустая строка между разделами

    doc.save(out_path)
    print(f'Документ сохранён: {out_path}')
    print(f'Разделов с формулами: {sum(1 for s in sections if s["formulas"])}')
    total = sum(len(s["formulas"]) for s in sections)
    print(f'Формул всего: {total}')


# ── Точка входа ──────────────────────────────────────────────────────────────

if __name__ == '__main__':
    sections = parse_md(MD_FILE)
    build_doc(sections, OUT_FILE)
