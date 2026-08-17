"""
build_full_doc.py
Конвертирует DST_ALGORITHM_FULL.md в полноценный Word-документ:
  - Весь текст (заголовки, абзацы, таблицы, списки, примечания)
  - LaTeX-блоки $$...$$ -> объект формулы Word (OMML) по центру
  - Inline-формулы (строки "> формула") -> объект формулы Word (OMML) по центру
  - При ошибке OMML -> текст выделяется красным

Зависимости: python-docx, latex2mathml, mathml2omml, lxml
"""

import re
from pathlib import Path
from lxml import etree
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import latex2mathml.converter as lc
import mathml2omml

MD_FILE  = Path(__file__).parent / "DST_ALGORITHM_FULL.md"
OUT_FILE = Path(__file__).parent / "DST_ALGORITHM_FULL.docx"

OMML_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

# ══════════════════════════════════════════════════════════════════
# Unicode -> LaTeX
# ══════════════════════════════════════════════════════════════════

# Порядок важен: сначала длинные строки
UNICODE_TO_LATEX = [
    # Каллиграфические множества (длинные сначала)
    ('\u210c_{P,D}',    r'\mathcal{E}_{P,D}'),
    ('\u210c_{P,F}',    r'\mathcal{E}_{P,F}'),
    ('\u210c_{P,U}',    r'\mathcal{E}_{P,U}'),
    ('\u210c',          r'\mathcal{E}'),
    ('\U0001d49e_{X,P}',r'\mathcal{C}_{X,P}'),
    ('\U0001d49e_all',  r'\mathcal{C}_{all}'),
    ('\U0001d49e',      r'\mathcal{C}'),
    ('\U0001d4b1_P',    r'\mathcal{V}_P'),
    ('\U0001d4b1_{P,D}',r'\mathcal{V}_{P,D}'),
    ('\U0001d4b1',      r'\mathcal{V}'),
    ('\U0001d4ae_C',    r'\mathcal{S}_C'),
    ('\U0001d4ae_F',    r'\mathcal{S}_F'),
    ('\U0001d4ae',      r'\mathcal{S}'),
    ('\U0001d4a2',      r'\mathcal{G}'),
    ('\U0001d4a6',      r'\mathcal{K}'),
    ('\U0001d4ab',      r'\mathcal{P}'),
    ('\U0001d4ac',      r'\mathcal{Q}'),
    ('\U0001d4b0',      r'\mathcal{U}'),
    ('\U0001d49f',      r'\mathcal{D}'),
    ('\u2131',          r'\mathcal{F}'),
    # Суммирование с нижним индексом
    ('\u2211_{',        r'\sum_{'),
    ('\u2211',          r'\sum '),
    # Логические операторы
    ('\u2208',          r'\in '),
    ('\u2209',          r'\notin '),
    ('\u2229',          r'\cap '),
    ('\u222a',          r'\cup '),
    ('\u2205',          r'\emptyset '),
    ('\u2260',          r'\neq '),
    ('\u2264',          r'\leq '),
    ('\u2265',          r'\geq '),
    ('\u2200',          r'\forall '),
    ('\u2203',          r'\exists '),
    ('\u2192',          r'\to '),
    ('\u2190',          r'\leftarrow '),
    # Математические символы
    ('\u2212',          r'-'),
    ('\u00b7',          r'\cdot '),
    ('\u00b1',          r'\pm '),
    ('\u207b\u00b9',    r'^{-1}'),
    # Греческие
    ('\u03ba',          r'\kappa '),
    ('\u03bb',          r'\lambda '),
    ('\u03c3',          r'\sigma '),
    ('\u03bc',          r'\mu '),
    ('\u03b4',          r'\delta '),
    ('\u03c1',          r'\rho '),
    ('\u03b3',          r'\gamma '),
    ('\u03c0',          r'\pi '),
    ('\u03c9',          r'\omega '),
    ('\u0394',          r'\Delta '),
    ('\u0398',          r'\Theta '),
    ('\u03c4',          r'\tau '),
    # Специальные обозначения
    ("m\u2032",         r"m'"),
    ("m*",              r"m^*"),
    ('\u2032',          r"'"),
    # Индексные и надстрочные символы
    ('\u1d62',          r'_i'),
    ('\u2C7C',          r'_j'),
    # Спецсимволы
    ('\u27e8',          r'\langle '),
    ('\u27e9',          r'\rangle '),
    # Ненужные escape
    (':=',              r':= '),
]


def unicode_to_latex(text):
    """Преобразует Unicode-формулу в LaTeX."""
    result = text
    # Убираем HTML-сущности
    result = result.replace('&nbsp;', ' ').replace('&amp;', '&')
    # Убираем Markdown bold/italic внутри формулы
    result = re.sub(r'\*\*(.+?)\*\*', r'\1', result)
    result = re.sub(r'\*(.+?)\*',    r'\1', result)
    # Убираем суррогаты (битые Unicode-символы)
    result = result.encode('utf-16', 'surrogatepass').decode('utf-16')
    # Нижние индексы-буквы Unicode (ᵢ, ⱼ)
    sub_map = {'i': 'i', 'j': 'j', 'e': 'e', 'o': 'o', 'a': 'a', 'x': 'x'}
    # Диапазон Unicode subscript letters: U+2090–U+209C, U+1D62–U+1D6A
    SUBSCRIPT_LETTERS = {
        '\u1d62': 'i', '\u1d63': 'r', '\u1d64': 'u', '\u1d65': 'v',
        '\u1d66': '\u03b2', '\u1d67': '\u03b3', '\u1d68': '\u03c1',
        '\u1d69': '\u03c6', '\u1d6a': '\u03c7',
        '\u2c7c': 'j',
        '\u2090': 'a', '\u2091': 'e', '\u2095': 'h', '\u2096': 'k',
        '\u2097': 'l', '\u2098': 'm', '\u2099': 'n', '\u209a': 'p',
        '\u209b': 's', '\u209c': 't',
    }
    for sub_ch, letter in SUBSCRIPT_LETTERS.items():
        result = result.replace(sub_ch, '_{' + letter + '}')
    # Основные замены
    for uni, lat in UNICODE_TO_LATEX:
        result = result.replace(uni, lat)
    # Логическое ИЛИ/И
    result = result.replace('\u2228', r'\lor ')
    result = result.replace('\u2227', r'\land ')
    # Составные имена переменных с подчёркиванием: raw_w -> \mathrm{raw\_w}
    result = re.sub(r'\braw_([a-zA-Z])',
                    lambda m: r'\mathrm{raw\_' + m.group(1) + '}',
                    result)
    # Двойное подчёркивание в именах переменных (s_c_eff -> s_{c\_eff})
    result = re.sub(r'([a-zA-Z])_([a-zA-Z]+)_([a-zA-Z]+)',
                    lambda m: m.group(1) + '_{' + m.group(2) + r'\_' + m.group(3) + '}',
                    result)
    # Заменяем оставшиеся не-ASCII символы на '?'
    safe = []
    for ch in result:
        if ord(ch) > 127:
            safe.append('?')
        else:
            safe.append(ch)
    return ''.join(safe)


def fix_omml_ns(omml_str):
    """Добавляет namespace к корневому элементу m:oMath."""
    ns_attr = 'xmlns:m="{}"'.format(OMML_NS)
    return omml_str.replace('<m:oMath>', '<m:oMath {}>'.format(ns_attr), 1)


def make_omml_element(latex_src):
    """LaTeX str -> lxml Element (OMML). Возвращает None при ошибке."""
    try:
        mml = lc.convert(latex_src)
        omml_str = mathml2omml.convert(mml)
        fixed = fix_omml_ns(omml_str)
        return etree.fromstring(fixed.encode('utf-8'))
    except Exception:
        return None


def add_formula_para(doc, latex_src, fallback_text, centered=True):
    """Добавляет абзац с OMML-формулой; при ошибке - текст."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    el = make_omml_element(latex_src)
    if el is not None:
        p._p.append(el)
        return True
    else:
        run = p.add_run(fallback_text)
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        return False


# ══════════════════════════════════════════════════════════════════
# Определение типа строки
# ══════════════════════════════════════════════════════════════════

MATH_MARKERS = [
    '=', '\u2260', '\u2205', '\u2211', '\u2208', '\u2209', '\u222a', '\u2229',
    '\u2264', '\u2265', '\u00b7', '\u2212', '|{', '}|', ':=',
    '\u03ba', '\u03bb', '\u03c3', '\u03bc', '\u03b4', '\u03c1', '\u03b3',
    '\u03c0', '\u03c9', '\u0394',
    'm(', "m'", 'm*', 'BetP', 'num(', 'BPA_', '+=',
    '\U0001d4ae', '\U0001d4b1', '\u210c', '\U0001d49e', '\u2131',
]

NOTE_STARTS = (
    '**\u0412\u0435\u0440\u0441\u0438\u044f',  # **Версия
    '**\u0420\u0435\u0430\u043b\u0438\u0437',   # **Реализ
    '**\u041a\u043b\u044e\u0447\u0435\u0432',   # **Ключев
    '\u041a\u043b\u044e\u0447\u0435\u0432',      # Ключев
    '\u0412 \u043d\u0430\u0441\u0442\u043e',    # В насто
    '\u0415\u0441\u043b\u0438 ',                 # Если
)


def is_formula_quote(line):
    """Строка-цитата > содержит формулу (не примечание)."""
    if not line.startswith('>'):
        return False
    content = line.lstrip('> ').strip()
    if not content:
        return False
    has_math = any(m in content for m in MATH_MARKERS)
    is_note = content.startswith('**') or any(content.startswith(s) for s in NOTE_STARTS)
    return has_math and not is_note


def strip_inline_md(text):
    """Убирает **bold**, *italic* маркеры, возвращает чистый текст."""
    t = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    t = re.sub(r'\*(.+?)\*', r'\1', t)
    return t


# ══════════════════════════════════════════════════════════════════
# Runs с inline-форматированием
# ══════════════════════════════════════════════════════════════════

def add_inline_runs(para, text, base_size=None):
    """Парсит **bold**, *italic*, `code` и добавляет runs в абзац."""
    parts = re.split(r'(\*\*[^*]+?\*\*|\*[^*]+?\*|`[^`]+?`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = para.add_run(part[2:-2])
            r.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            r = para.add_run(part[1:-1])
            r.italic = True
        elif part.startswith('`') and part.endswith('`'):
            r = para.add_run(part[1:-1])
            r.font.name = 'Courier New'
            r.font.size = Pt(9)
        else:
            para.add_run(part)
    if base_size:
        for run in para.runs:
            if not run.font.size:
                run.font.size = base_size


# ══════════════════════════════════════════════════════════════════
# Основной цикл построения документа
# ══════════════════════════════════════════════════════════════════

def build_doc(md_path, out_path):
    doc = Document()

    sec = doc.sections[0]
    sec.top_margin    = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2)

    lines = md_path.read_text(encoding='utf-8').splitlines()
    i = 0
    n = len(lines)
    ok_count = 0
    err_count = 0
    err_list = []

    while i < n:
        line = lines[i]

        # Пустая строка
        if not line.strip():
            i += 1
            continue

        # Горизонтальная линия ---
        if re.match(r'^-{3,}$', line.strip()):
            doc.add_paragraph()
            i += 1
            continue

        # Заголовок #, ##, ###, ####
        m = re.match(r'^(#{1,4})\s+(.+)', line)
        if m:
            lvl = min(len(m.group(1)), 4)
            text = m.group(2).strip()
            p = doc.add_heading(text, level=lvl)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
            continue

        # LaTeX-блок $$
        if line.strip() == '$$':
            latex_lines = []
            i += 1
            while i < n and lines[i].strip() != '$$':
                latex_lines.append(lines[i].strip())
                i += 1
            i += 1
            latex_src = ' '.join(latex_lines)
            ok = add_formula_para(doc, latex_src, latex_src, centered=True)
            ok_count += ok
            if not ok:
                err_count += 1
                err_list.append(latex_src[:60])
            continue

        # Цитаты > (формулы и примечания)
        if line.startswith('>'):
            quote_lines = []
            while i < n and lines[i].startswith('>'):
                quote_lines.append(lines[i])
                i += 1
            for ql in quote_lines:
                content = re.sub(r'^>\s*', '', ql).strip()
                if not content:
                    continue
                if is_formula_quote(ql):
                    latex_src = unicode_to_latex(content)
                    ok = add_formula_para(doc, latex_src, content, centered=True)
                    ok_count += ok
                    if not ok:
                        err_count += 1
                        err_list.append(content[:60])
                else:
                    # Примечание — курсивный абзац
                    p = doc.add_paragraph()
                    r = p.add_run(strip_inline_md(content))
                    r.italic = True
                    r.font.size = Pt(10)
            continue

        # Markdown-таблица |col|col|
        if line.startswith('|'):
            tbl_lines = []
            while i < n and lines[i].startswith('|'):
                tbl_lines.append(lines[i])
                i += 1
            rows = []
            for tl in tbl_lines:
                cells = [c.strip() for c in tl.split('|')]
                cells = [c for c in cells if c]
                if all(re.match(r'^[-:]+$', c) for c in cells):
                    continue
                rows.append(cells)
            if rows:
                ncols = max(len(r) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=ncols)
                tbl.style = 'Table Grid'
                for ri, row_data in enumerate(rows):
                    for ci, val in enumerate(row_data[:ncols]):
                        cell = tbl.rows[ri].cells[ci]
                        cell.text = ''
                        add_inline_runs(cell.paragraphs[0], val, base_size=Pt(9))
                        if ri == 0:
                            for run in cell.paragraphs[0].runs:
                                run.bold = True
                doc.add_paragraph()
            continue

        # Маркированный список - item
        if re.match(r'^[-*]\s', line):
            list_lines = []
            while i < n and re.match(r'^[-*]\s', lines[i]):
                list_lines.append(lines[i])
                i += 1
            for ll in list_lines:
                text = re.sub(r'^[-*]\s+', '', ll)
                p = doc.add_paragraph(style='List Bullet')
                add_inline_runs(p, text)
            continue

        # Нумерованный список 1. item
        if re.match(r'^\d+\.\s', line):
            list_lines = []
            while i < n and re.match(r'^\d+\.\s', lines[i]):
                list_lines.append(lines[i])
                i += 1
            for ll in list_lines:
                text = re.sub(r'^\d+\.\s+', '', ll)
                p = doc.add_paragraph(style='List Number')
                add_inline_runs(p, text)
            continue

        # Code block ```
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1
            if code_lines:
                p = doc.add_paragraph()
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'),   'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'),  'F5F5F5')
                pPr.append(shd)
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(8.5)
            continue

        # Курсивный подзаголовок *Уровень 0...*
        if re.match(r'^\*[^*]', line) and line.rstrip().endswith('*'):
            text = line.strip('* ')
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.bold = True
            r.italic = True
            r.font.size = Pt(10.5)
            i += 1
            continue

        # Обычный абзац
        p = doc.add_paragraph(style='Normal')
        add_inline_runs(p, line)
        i += 1

    doc.save(out_path)
    total = ok_count + err_count
    print('Saved: ' + str(out_path))
    print('Formulas total: ' + str(total) + '  OK: ' + str(ok_count) + '  ERR: ' + str(err_count))
    if err_list:
        print('Failed formulas:')
        for e in err_list[:10]:
            print('  ' + e)


if __name__ == '__main__':
    build_doc(MD_FILE, OUT_FILE)
